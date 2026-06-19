#!/usr/bin/env python3
"""Generate Word document: model metrics, depth, and algorithm specifications."""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "Model_Metrics_and_Depth.docx"
PLOTS = DOCS / "metrics_plots"
PLOTS.mkdir(parents=True, exist_ok=True)
METRICS_JSON = DOCS / "model_metrics_snapshot.json"

STUDENT = "Nick Einstein"
TITLE = "Model Metrics, Algorithm Depth & Evaluation"


def load_metrics() -> dict:
    if METRICS_JSON.exists():
        return json.loads(METRICS_JSON.read_text())
    return {}


def make_charts(m: dict) -> dict[str, Path]:
    paths: dict[str, Path] = {}
    dark, face = "#06080f", "#0c1019"

    def style(ax, fig):
        ax.set_facecolor(face)
        fig.patch.set_facecolor(dark)
        ax.tick_params(colors="white")
        for label in (ax.xaxis.label, ax.yaxis.label, ax.title):
            if label:
                label.set_color("white")
        ax.title.set_color("white")
        for spine in ax.spines.values():
            spine.set_color("#333")

    # Metrics bar chart
    fig, ax = plt.subplots(figsize=(8, 4.5))
    names = ["ROC-AUC", "Avg Precision", "F1", "Precision", "Recall", "Balanced Acc."]
    vals = [
        m.get("roc_auc", 0),
        m.get("average_precision", 0),
        m.get("f1", 0),
        m.get("precision", 0),
        m.get("recall", 0),
        m.get("balanced_accuracy", 0),
    ]
    bars = ax.bar(names, vals, color=["#3dffa8", "#5ec8ff", "#ffc857", "#ff9f43", "#ff4d6d", "#a78bfa"])
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Score")
    ax.set_title("Classification Metrics (Test Set, threshold = 0.5)")
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, v + 0.02, f"{v:.3f}", ha="center", color="white", fontsize=9)
    style(ax, fig)
    p = PLOTS / "metrics_bar.png"
    fig.tight_layout()
    fig.savefig(p, dpi=150, facecolor=fig.get_facecolor())
    plt.close(fig)
    paths["metrics_bar"] = p

    # Confusion matrix
    cm = m.get("confusion_matrix", [[99484, 114], [131, 271]])
    fig, ax = plt.subplots(figsize=(5, 4))
    im = ax.imshow(cm, cmap="Blues")
    ax.set_xticks([0, 1], labels=["Pred Legit", "Pred Fraud"])
    ax.set_yticks([0, 1], labels=["Actual Legit", "Actual Fraud"])
    for i in range(2):
        for j in range(2):
            ax.text(j, i, f"{cm[i][j]:,}", ha="center", va="center", color="white", fontsize=14, fontweight="bold")
    ax.set_title("Confusion Matrix (n=100,000 test rows)")
    style(ax, fig)
    p = PLOTS / "confusion_matrix.png"
    fig.tight_layout()
    fig.savefig(p, dpi=150, facecolor=fig.get_facecolor())
    plt.close(fig)
    paths["confusion"] = p

    # Threshold sweep
    sweep = m.get("threshold_sweep", {})
    if sweep:
        fig, ax = plt.subplots(figsize=(7, 4))
        ths = sorted(float(k) for k in sweep.keys())
        prec = [sweep[str(t)]["precision"] for t in ths]
        rec = [sweep[str(t)]["recall"] for t in ths]
        f1s = [sweep[str(t)]["f1"] for t in ths]
        ax.plot(ths, prec, "o-", label="Precision", color="#5ec8ff")
        ax.plot(ths, rec, "s-", label="Recall", color="#ff4d6d")
        ax.plot(ths, f1s, "^-", label="F1", color="#3dffa8")
        ax.axvline(0.5, color="#888", linestyle="--", label="Default threshold")
        ax.set_xlabel("Classification threshold")
        ax.set_ylabel("Score")
        ax.set_title("Metrics vs Fraud Probability Threshold")
        ax.legend(facecolor=face, edgecolor="white", labelcolor="white")
        style(ax, fig)
        p = PLOTS / "threshold_sweep.png"
        fig.tight_layout()
        fig.savefig(p, dpi=150, facecolor=fig.get_facecolor())
        plt.close(fig)
        paths["threshold"] = p

    # Feature importance
    fi_path = ROOT / "artifacts" / "feature_importance.csv"
    if fi_path.exists():
        df = pd.read_csv(fi_path).head(13)
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.barh(df["feature"][::-1], df["importance"][::-1], color="#3dffa8")
        ax.set_xlabel("CatBoost Feature Importance")
        ax.set_title("Feature Importance (13 inputs)")
        style(ax, fig)
        p = PLOTS / "feature_importance.png"
        fig.tight_layout()
        fig.savefig(p, dpi=150, facecolor=fig.get_facecolor())
        plt.close(fig)
        paths["fi"] = p

    # Model stack diagram as text - skip

    return paths


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        p = doc.add_paragraph(caption)
        p.runs[0].italic = True
        doc.add_picture(str(path), width=Inches(5.8))


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def build() -> Path:
    m = load_metrics()
    if not m:
        raise SystemExit("Run metrics computation first — see scripts/compute_model_metrics.py")

    charts = make_charts(m)
    cm = m.get("confusion_matrix", [[0, 0], [0, 0]])
    tn, fp, fn, tp = cm[0][0], cm[0][1], cm[1][0], cm[1][1]

    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{STUDENT}\nSentinel Fraud Detection Project").font.size = Pt(12)
    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. Algorithms in the Stack")
    doc.add_paragraph(
        "The Sentinel platform uses three complementary algorithms: (1) CatBoost gradient boosted "
        "decision trees for supervised fraud classification; (2) SMOTE for training-time class "
        "balancing; (3) Isolation Forest for unsupervised anomaly scoring in the real-time stream. "
        "A fourth analytical method — Kolmogorov–Smirnov drift testing — monitors feature distributions "
        "but is not a predictive model."
    )
    add_table(
        doc,
        ["Algorithm", "Type", "Role", "Artifact"],
        [
            ["CatBoost", "Supervised GBDT", "Primary fraud probability", "artifacts/models/catboost_fraud.cbm"],
            ["SMOTE", "Resampling", "Balance training labels only", "Applied in-memory at train time"],
            ["StandardScaler", "Preprocessing", "Normalize features before trees", "artifacts/scaler.joblib"],
            ["Isolation Forest", "Unsupervised ensemble", "Anomaly score in live stream", "artifacts/models/anomaly_iforest.joblib"],
            ["KS test", "Statistical test", "Drift vs training baseline", "artifacts/baselines/feature_distributions.joblib"],
        ],
    )

    # 2. CatBoost depth
    add_heading(doc, "2. CatBoost Classifier — Model Depth & Structure")
    doc.add_paragraph(
        "CatBoost (Categorical Boosting) builds an ensemble of decision trees sequentially. Each new "
        "tree corrects errors from prior trees using gradient boosting. Our configuration targets "
        "tabular fraud data with severe class imbalance."
    )

    add_heading(doc, "2.1 Hyperparameters (config.yaml)", level=2)
    depth = m.get("catboost_depth", 6)
    iters = m.get("catboost_iterations", 500)
    tree_count = m.get("catboost_tree_count", iters)
    add_table(
        doc,
        ["Parameter", "Value", "Meaning"],
        [
            ["depth", str(depth), "Maximum depth of each decision tree (levels from root to leaf)"],
            ["iterations", str(iters), "Number of boosting trees built in the ensemble"],
            ["learning_rate", "0.05", "Shrinkage per tree — smaller steps, less overfit"],
            ["eval_metric", "AUC", "Optimization target during training (ROC area)"],
            ["auto_class_weights", "Balanced", "Penalize misclassified minority (fraud) class more"],
            ["random_seed", "42", "Reproducible training"],
            ["verbose", "100", "Log every 100 iterations"],
        ],
    )

    add_heading(doc, "2.2 What “Depth = 6” Means", level=2)
    doc.add_paragraph(
        "Each tree can ask up to six sequential splitting questions (e.g. category < X, then amt < Y, …) "
        "before reaching a leaf prediction. Depth 6 balances complexity and generalization: deep enough "
        "to capture interactions (category × amount × time), shallow enough to avoid memorizing rare "
        "fraud patterns on SMOTE-synthetic points."
    )
    add_bullets(
        doc,
        [
            f"Trees in final model: {tree_count} (one per boosting iteration).",
            f"Input features per tree: {m.get('feature_count', 13)} (all scaled numeric columns).",
            f"Training rows after SMOTE: {m.get('train_rows_after_smote', '—'):,} (balanced classes).",
            "Output: fraud_probability ∈ [0, 1] via sigmoid of ensemble logit sum.",
            "Decision rule at serve time: fraud_prediction = 1 if probability ≥ 0.5.",
        ],
    )

    add_heading(doc, "2.3 Mathematical Form (Conceptual)", level=2)
    add_code(
        doc,
        """F(x) = Σ_{t=1}^{T} η · h_t(x)     where T = 500 trees, η = 0.05 learning rate
P(fraud) = σ(F(x))                  σ = logistic sigmoid
h_t = decision tree of depth ≤ 6    each h_t fits residual gradients""",
    )

    add_heading(doc, "2.4 Why CatBoost for This Problem", level=2)
    add_bullets(
        doc,
        [
            "Handles mixed numeric features after encoding without manual one-hot explosion.",
            "Ordered boosting reduces prediction shift on small fraud samples.",
            "auto_class_weights complements SMOTE for minority-class recall.",
            "Strong ROC-AUC on tabular financial data vs single decision trees or logistic regression.",
        ],
    )

    doc.add_page_break()

    # 3. SMOTE & preprocessing depth
    add_heading(doc, "3. SMOTE & Preprocessing Depth")
    add_table(
        doc,
        ["Component", "Setting", "Detail"],
        [
            ["SMOTE k_neighbors", "5", "Synthetic fraud points interpolated from 5 nearest real fraud neighbors"],
            ["SMOTE sampling_strategy", "auto", "Resample minority until balanced with majority class"],
            ["SMOTE applied to", "Train only", "Test set keeps natural ~0.4% fraud rate"],
            ["StandardScaler", "13 features", "μ=0, σ=1 per column; fit on SMOTE-balanced train"],
            ["PCA", "2 components", "Visualization only (pca_train.png); NOT fed to CatBoost"],
        ],
    )

    # 4. Isolation Forest
    add_heading(doc, "4. Isolation Forest — Anomaly Model Depth")
    add_table(
        doc,
        ["Parameter", "Value", "Meaning"],
        [
            ["n_estimators", "100", "Number of random trees in the forest"],
            ["contamination", "0.002", "Expected anomaly proportion (~0.2% of training sample)"],
            ["Tree structure", "Random splits", "Each tree isolates points; shorter path = more anomalous"],
            ["max_samples", "auto (sklearn default)", "Subsample per tree for diversity"],
            ["Output", "decision_function", "Inverted to anomaly_score (higher = more unusual)"],
        ],
    )
    doc.add_paragraph(
        "Isolation Forest does not predict fraud labels directly. It flags statistically rare "
        "transactions in the real-time WebSocket stream, complementing CatBoost when attack patterns "
        "are novel and not present in labelled training data."
    )

    doc.add_page_break()

    # 5. Metrics definitions
    add_heading(doc, "5. Evaluation Metrics — Definitions")
    add_table(
        doc,
        ["Metric", "Formula / Definition", "Why we use it"],
        [
            ["ROC-AUC", "Area under TPR vs FPR curve", "Ranking quality; robust to imbalance"],
            ["Average Precision", "Area under precision-recall curve", "Focuses on positive (fraud) class"],
            ["Precision", "TP / (TP + FP)", "Of flagged transactions, how many are truly fraud"],
            ["Recall", "TP / (TP + FN)", "Of all fraud, how many we catch"],
            ["F1", "2·P·R / (P + R)", "Harmonic mean of precision and recall"],
            ["Accuracy", "(TP+TN) / N", "Misleading when fraud < 1% — reported for completeness"],
            ["Balanced accuracy", "Mean of per-class recall", "Fairer than raw accuracy under imbalance"],
        ],
    )

    # 6. Results
    add_heading(doc, "6. Measured Results (fraudTest hold-out)")
    doc.add_paragraph(
        f"Evaluation on {m.get('test_rows', 100000):,} test rows "
        f"({m.get('test_fraud_count', 402)} fraud, rate {m.get('test_fraud_rate', 0.004)*100:.2f}%). "
        "Threshold = 0.5 unless noted."
    )

    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["ROC-AUC", f"{m.get('roc_auc', 0):.4f}"],
            ["Average Precision (PR-AUC)", f"{m.get('average_precision', 0):.4f}"],
            ["F1 Score", f"{m.get('f1', 0):.4f}"],
            ["Precision", f"{m.get('precision', 0):.4f}"],
            ["Recall", f"{m.get('recall', 0):.4f}"],
            ["Accuracy", f"{m.get('accuracy', 0):.4f}"],
            ["Balanced Accuracy", f"{m.get('balanced_accuracy', 0):.4f}"],
        ],
    )

    add_heading(doc, "6.1 Confusion Matrix", level=2)
    add_table(
        doc,
        ["", "Predicted Legit", "Predicted Fraud"],
        [
            ["Actual Legit", f"{tn:,} (True Negative)", f"{fp:,} (False Positive)"],
            ["Actual Fraud", f"{fn:,} (False Negative)", f"{tp:,} (True Positive)"],
        ],
    )
    doc.add_paragraph(
        f"Interpretation: {tp} fraud cases correctly flagged; {fn} fraud cases missed; "
        f"{fp} legitimate transactions incorrectly flagged as fraud."
    )
    add_figure(doc, charts.get("confusion", Path()), "Figure 1: Confusion matrix heatmap.")

    add_heading(doc, "6.2 Per-Class Classification Report", level=2)
    if m.get("classification_report"):
        add_code(doc, m["classification_report"].strip())

    add_figure(doc, charts.get("metrics_bar", Path()), "Figure 2: Summary of key metrics at threshold 0.5.")

    add_heading(doc, "6.3 Threshold Sensitivity", level=2)
    doc.add_paragraph(
        "Lowering the threshold increases recall (catch more fraud) but adds false positives. "
        "Raising it improves precision but misses more fraud."
    )
    sweep = m.get("threshold_sweep", {})
    if sweep:
        rows = []
        for th in sorted(sweep.keys(), key=float):
            s = sweep[th]
            rows.append([th, f"{s['precision']:.3f}", f"{s['recall']:.3f}", f"{s['f1']:.3f}", str(s["flagged"])])
        add_table(doc, ["Threshold", "Precision", "Recall", "F1", "Transactions flagged"], rows)
    add_figure(doc, charts.get("threshold", Path()), "Figure 3: Precision, recall, and F1 vs threshold.")

    doc.add_page_break()

    # 7. Feature importance
    add_heading(doc, "7. Feature Importance (Model Interpretation)")
    doc.add_paragraph(
        "CatBoost reports feature importance as the total gain from splits using each variable "
        "across all 500 trees. Higher values mean the feature contributed more to reducing loss."
    )
    fi = m.get("feature_importance", [])
    if fi:
        add_table(
            doc,
            ["Rank", "Feature", "Importance", "Role"],
            [
                [
                    str(i + 1),
                    row["feature"],
                    f"{row['importance']:.2f}",
                    _feature_role(row["feature"]),
                ]
                for i, row in enumerate(fi)
            ],
        )
    add_figure(doc, charts.get("fi", Path()), "Figure 4: CatBoost feature importance.")

    # 8. Drift & governance metrics
    add_heading(doc, "8. Post-Model Analysis Metrics")
    add_bullets(
        doc,
        [
            "KS statistic per feature: max gap between train and test cumulative distributions.",
            "Drift detected if p-value < α = 0.05 (config monitoring.drift.alpha).",
            "Governance min_roc_auc = 0.85 — model must exceed this to avoid retrain recommendation.",
            "Production model version: v_f381139c (registry artifacts/registry/registry.json).",
        ],
    )

    add_heading(doc, "9. Summary")
    doc.add_paragraph(
        f"The primary CatBoost ensemble uses {tree_count} trees of depth {depth}, learning rate 0.05, "
        f"and balanced class weights on {m.get('feature_count', 13)} scaled features. On the held-out "
        f"test set it achieves ROC-AUC {m.get('roc_auc', 0):.4f} with recall "
        f"{m.get('recall', 0):.2%} and precision {m.get('precision', 0):.2%} at the default 0.5 "
        "threshold. High ROC-AUC confirms strong ranking; moderate recall reflects the difficulty "
        "of fraud detection under extreme imbalance and adversarial behaviour."
    )

    foot = doc.add_paragraph(f"Prepared by {STUDENT}")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT))
    return OUT


def _feature_role(name: str) -> str:
    roles = {
        "category": "Merchant type — dominant predictor",
        "amt": "Transaction amount",
        "gender": "Cardholder gender (encoded)",
        "city_pop": "City population",
        "trans_hour": "Hour of transaction",
        "unix_time": "Epoch timestamp",
        "state": "US state (encoded)",
        "lat": "Cardholder latitude",
        "long": "Cardholder longitude",
        "trans_dow": "Day of week",
        "merch_long": "Merchant longitude",
        "distance_km": "Haversine distance km",
        "merch_lat": "Merchant latitude",
    }
    return roles.get(name, "Model input")


if __name__ == "__main__":
    path = build()
    print(f"Model metrics document saved: {path}")
